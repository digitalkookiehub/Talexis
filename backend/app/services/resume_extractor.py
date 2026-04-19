"""Resume text extraction supporting PDF, DOCX, and plain text."""
import logging
import os
from io import BytesIO

logger = logging.getLogger(__name__)


def extract_text_from_file(filepath: str) -> str:
    """Extract text from a resume file. Supports PDF, DOCX, TXT."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(filepath)
    if ext in (".docx", ".doc"):
        return _extract_docx(filepath)
    if ext == ".txt":
        return _extract_txt(filepath)

    # Fallback: try as plain text
    return _extract_txt(filepath)


def _extract_pdf(filepath: str) -> str:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        text_parts = []
        for page in reader.pages:
            try:
                text_parts.append(page.extract_text() or "")
            except Exception as e:
                logger.warning("Failed to extract page text: %s", str(e))
        text = "\n".join(text_parts).strip()
        if not text:
            logger.warning("PDF extraction returned empty text for %s", filepath)
            return "Could not extract text from PDF (may be scanned image)."
        return text
    except ImportError:
        logger.error("pypdf not installed")
        return "PDF parsing not available (pypdf missing)."
    except Exception as e:
        logger.error("PDF extraction failed: %s", str(e))
        return f"PDF extraction error: {str(e)}"


def _extract_docx_full(doc_source: object) -> str:
    """Extract ALL text from a DOCX document — paragraphs, tables, headers, footers, and text boxes.

    Many resumes use table-based layouts where all content is inside table cells,
    not in top-level paragraphs. This function handles that.
    """
    from docx import Document
    from docx.oxml.ns import qn

    if isinstance(doc_source, (str, bytes, os.PathLike)):
        doc = Document(doc_source)
    else:
        doc = Document(doc_source)

    text_parts: list[str] = []

    # 1. Top-level paragraphs
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            text_parts.append(t)

    # 2. Tables — iterate all rows/cells (resumes often use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            row_texts: list[str] = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                text_parts.append("\n".join(row_texts))

    # 3. Headers and footers (some resumes put name/contact in header)
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header and header.is_linked_to_previous is False:
                for p in header.paragraphs:
                    t = p.text.strip()
                    if t:
                        text_parts.append(t)
        for footer in [section.footer, section.first_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for p in footer.paragraphs:
                    t = p.text.strip()
                    if t:
                        text_parts.append(t)

    # 4. Text boxes (floating content in shapes/text frames)
    try:
        body = doc.element.body
        for txbx in body.iter(qn("w:txbxContent")):
            for p in txbx.iter(qn("w:p")):
                runs = p.findall(qn("w:r"))
                p_text = "".join(
                    r.find(qn("w:t")).text or ""
                    for r in runs
                    if r.find(qn("w:t")) is not None
                )
                if p_text.strip():
                    text_parts.append(p_text.strip())
    except Exception as e:
        logger.debug("Text box extraction skipped: %s", str(e))

    return "\n".join(text_parts)


def _extract_docx(filepath: str) -> str:
    """Extract text from DOCX using python-docx — handles tables, text boxes, headers."""
    try:
        text = _extract_docx_full(filepath)
        if text.strip():
            return text
        logger.warning("DOCX extraction returned empty text for %s", filepath)
        return ""
    except ImportError:
        logger.error("python-docx not installed")
        return "DOCX parsing not available."
    except Exception as e:
        logger.error("DOCX extraction failed: %s", str(e))
        return f"DOCX extraction error: {str(e)}"


def _extract_txt(filepath: str) -> str:
    """Extract text from plain text file."""
    try:
        with open(filepath, "rb") as f:
            content = f.read()
        # Try UTF-8 first, fall back to latin-1
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1", errors="ignore")
    except Exception as e:
        logger.error("Text extraction failed: %s", str(e))
        return ""


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """Extract text from raw bytes (used for in-memory parsing)."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        except Exception as e:
            logger.error("PDF byte extraction failed: %s", str(e))
            return ""

    if ext in (".docx", ".doc"):
        try:
            text = _extract_docx_full(BytesIO(content))
            return text
        except Exception as e:
            logger.error("DOCX byte extraction failed: %s", str(e))
            return ""

    # Plain text
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1", errors="ignore")
